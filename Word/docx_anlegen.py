# docx_anlegen.py
import docx

d = docx.Document()

d.add_paragraph('Hallo, dies ist ein Absatz.')
d.add_paragraph('Und dies noch ein Absatz.')

p = d.paragraphs[1]
p.add_run('Ein besonders schöner Absatz!')

p.runs[1].bold = True

p = d.paragraphs[0]
p.style = 'Heading 2'

d.save('test22.docx')