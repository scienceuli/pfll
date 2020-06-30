# index2paragraph.py

"""
loads a docx document with Codes in front of every paragraph.
index_dict maps these codes (keys) to index entries (values).
At the end of each para a XE field with value and bookmark with key are added.
At the end of document hyperlinks are added, from value to key
"""

import docx

from docx import Document
from lxml import etree

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import re

index_dict = {
    "123": "erster Indexeintrag",
    "456": "zweiter Indexeintrag",
    "789": "dritter Indexeintrag",
}


def mark_index_entry(entry, paragraph):
    """
    adds XE field with entry at the end of paragraph
    """
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' XE "%s" ' % (entry)
    r.append(instrText)

    # hyperlink in docx xml
    # <w:hyperlink w:anchor="marke1" w:history="1"><w:r w:rsidRPr="00400BB4"><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>marke1</w:t></w:r></w:hyperlink>

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)


def add_bookmark(mark, paragraph):
    """
    adds bookmark with anchor = mark at the end of paragraph
    """
    run = paragraph.add_run()
    tag = run._r
    bookmarkStart = OxmlElement("w:bookmarkStart")
    bookmarkStart.set(qn("w:id"), "0")
    bookmarkStart.set(qn("w:name"), mark)
    tag.append(bookmarkStart)

    text = OxmlElement("w:r")
    tag.append(text)

    bookmarkEnd = OxmlElement("w:bookmarkEnd")
    bookmarkEnd.set(qn("w:id"), "0")
    bookmarkEnd.set(qn("w:name"), mark)
    tag.append(bookmarkEnd)

    return run


def add_hyperlink(paragraph, link_to, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param link_to: internal bookmark
    :param text: The text displayed for the link
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    hyperlink = OxmlElement("w:hyperlink")

    # Create the w:hyperlink tag and add needed values
    hyperlink.set(
        qn("w:anchor"), link_to,
    )

    # Create a w:r element
    new_run = OxmlElement("w:r")

    # Create a new w:rPr element
    rPr = OxmlElement("w:rPr")

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


doc = Document("index2para.docx")

for p in doc.paragraphs:
    for key, value in index_dict.items():
        if key in p.text:
            mark_index_entry(value, p)
            add_bookmark(key, p)

# add  hyperlinks with the index_dict
# add new paragraph for every link
for key, value in index_dict.items():
    p = doc.add_paragraph()
    hyperlink = add_hyperlink(p, key, value)


doc.save("index2para_bearb.docx")
