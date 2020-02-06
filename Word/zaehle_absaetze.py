import docx

doc = docx.Document('test_01.docx')

docout = docx.Document()

liste_der_absaetze = doc.paragraphs
anzahl_absaetze = len(liste_der_absaetze)
erster_absatz = liste_der_absaetze[0]
text_erster_absatz = erster_absatz.text

print(f"Das Dokument hat {anzahl_absaetze} Absätze.")
print(f"Der erste Absatz lautet: {text_erster_absatz}")


docout.add_paragraph("hallo\n\n\n")
docout.add_paragraph(text_erster_absatz.replace("Absatz", "Elefant"))
docout.add_paragraph(str(5*3))
docout.save('test_01_out.docx')


