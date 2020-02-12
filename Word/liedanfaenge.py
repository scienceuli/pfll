# liedanfaenge.py

import docx
import re

doc = docx.Document("Xire_Orixa-Satz-06.docx")

for para in doc.paragraphs:
    # print(para.style.name)
    r = re.match('Heading 3',para.style.name)

    if r: 
        t = para.text
        # print(t)
        new_para = doc.add_paragraph(t.strip())
        new_para.style = "Standard_uk"
        p = para._p
        p.addnext(new_para._p)

doc.save("Xire_Orixa-Satz-06_bearb.docx")

